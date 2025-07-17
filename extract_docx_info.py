#!/usr/bin/env python3
"""
Extract basic information from a docx file and output in JSON format.
"""

import json
import re
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Any

try:
    from docx import Document
    from docx.text.paragraph import Paragraph
    from docx.table import Table
    from docx.oxml.text.paragraph import CT_P
    from docx.oxml.table import CT_Tbl
    from docxlatex.docxlatex import Document as DocxLatexDocument
except ImportError as e:
    print(f"Please install required packages: {e}")
    print("Run: pip install python-docx docxlatex")
    exit(1)


class DocxInfoExtractor:
    def __init__(self, docx_path: str):
        self.docx_path = Path(docx_path)
        self.doc = Document(docx_path)
        self.current_year = datetime.now().year
        self.latex_content = None  # Cache for LaTeX conversion
        
    def extract_basic_info(self) -> Dict[str, Any]:
        """Extract basic information from the docx file"""
        text_content = self._get_full_text()
        
        basic_info = {
            "title": self._extract_title(text_content),
            "author": self._extract_author(text_content),
            "school": self._extract_school(text_content),
            "advisor": self._extract_advisor(text_content),
            "keywords": self._extract_keywords(text_content)
        }
        
        overall_stats = self._extract_overall_stats()
        chapter_stats = self._extract_chapter_stats()
        reference_stats = self._extract_reference_stats(text_content)
        
        return {
            "basic_info": basic_info,
            "overall_stats": overall_stats,
            "chapter_stats": chapter_stats,
            "reference_stats": reference_stats
        }
    
    def _get_full_text(self) -> str:
        """Get full text content from document"""
        full_text = []
        for element in self.doc.element.body:
            if isinstance(element, CT_P):
                paragraph = Paragraph(element, self.doc)
                full_text.append(paragraph.text)
            elif isinstance(element, CT_Tbl):
                table = Table(element, self.doc)
                for row in table.rows:
                    for cell in row.cells:
                        full_text.append(cell.text)
        return '\n'.join(full_text)
    
    def _extract_title(self, text: str) -> str:
        """Extract title from document"""
        # Look for title pattern starting with "题目:"
        pattern = r'题目[：:]?\s*([^\n]+)'
        match = re.search(pattern, text)
        if match:
            return match.group(1).strip()
        return "未识别"
    
    def _extract_author(self, text: str) -> str:
        """Extract author from document"""
        patterns = [
            r'学\s*生[：:]?\s*([^\s\n：:]+)',
            r'作\s*者[：:]?\s*([^\s\n：:]+)',
            r'姓\s*名[：:]?\s*([^\s\n：:]+)',
        ]
        
        for pattern in patterns:
            match = re.search(pattern, text)
            if match:
                author = match.group(1).strip()
                # Filter out long text that's clearly not a name
                if len(author) < 20 and not re.search(r'[。，！？]', author):
                    return author
        return "未识别"
    
    def _extract_school(self, text: str) -> str:
        """Extract school/department from document"""
        patterns = [
            r'学\s*院[：:]?\s*([^\s\n]+)',
            r'院\s*系[：:]?\s*([^\s\n]+)',
            r'系\s*别[：:]?\s*([^\s\n]+)',
        ]
        
        for pattern in patterns:
            match = re.search(pattern, text)
            if match:
                return match.group(1).strip()
        return "未识别"
    
    def _extract_advisor(self, text: str) -> str:
        """Extract advisor from document"""
        patterns = [
            r'指导教师[：:]?\s*([^\s\n]+)',
            r'导\s*师[：:]?\s*([^\s\n]+)',
            r'指导老师[：:]?\s*([^\s\n]+)',
        ]
        
        for pattern in patterns:
            match = re.search(pattern, text)
            if match:
                return match.group(1).strip()
        return "未识别"
    
    def _extract_keywords(self, text: str) -> List[str]:
        """Extract keywords from document"""
        pattern = r'关键词[：:]?\s*([^\n]+)'
        match = re.search(pattern, text)
        if match:
            keywords_text = match.group(1).strip()
            # Split by common delimiters
            keywords = re.split(r'[；;，,、\s]+', keywords_text)
            return [kw.strip() for kw in keywords if kw.strip()]
        return []
    
    def _extract_overall_stats(self) -> Dict[str, int]:
        """Extract overall statistics"""
        total_words = 0
        total_equations = 0
        total_paragraphs = 0
        total_images = 0
        total_tables = 0
        
        # Count paragraphs and words (Chinese + English)
        for paragraph in self.doc.paragraphs:
            if paragraph.text.strip():
                total_paragraphs += 1
                # Count Chinese characters and English words
                text = paragraph.text
                # Count Chinese characters
                chinese_chars = len(re.findall(r'[\u4e00-\u9fff]', text))
                # Count English words (split by whitespace, filter out Chinese)
                english_text = re.sub(r'[\u4e00-\u9fff]', ' ', text)
                english_words = len([word for word in english_text.split() if word.strip()])
                total_words += chinese_chars + english_words
        
        # Count tables
        total_tables = len(self.doc.tables)
        
        # Count images (inline shapes)
        for paragraph in self.doc.paragraphs:
            for run in paragraph.runs:
                if run.element.xpath('.//a:blip'):
                    total_images += 1
        
        # Count display equations using LaTeX conversion
        total_equations = self._count_display_equations()
        
        return {
            "total_words": total_words,
            "total_equations": total_equations,
            "total_paragraphs": total_paragraphs,
            "total_images": total_images,
            "total_tables": total_tables
        }
    
    def _extract_chapter_stats(self) -> Dict[str, Any]:
        """Extract chapter-based statistics"""
        chapters = []
        word_counts = []
        equation_counts = []
        table_counts = []
        image_counts = []
        paragraph_counts = []
        
        # Common chapter patterns - more specific to avoid TOC interference
        chapter_patterns = [
            (r'^(摘\s*要)$', '摘要'),
            (r'^(ABSTRACT)$', 'Abstract'),
            (r'^第(一|二|三|四|五|六|七|八|九|十)章\s+(.+)', lambda m: f'第{m.group(1)}章 {m.group(2)}'),
            (r'^第(一|二|三|四|五|六|七|八|九|十)章$', lambda m: f'第{m.group(1)}章'),
            (r'^(参考文献)$', '参考文献'),
        ]
        
        # Get total counts for distribution
        total_tables = len(self.doc.tables)
        total_images = 0
        for paragraph in self.doc.paragraphs:
            for run in paragraph.runs:
                if run.element.xpath('.//a:blip'):
                    total_images += 1
        
        current_chapter = None
        current_chapter_text = []
        current_paragraph_count = 0
        chapter_paragraphs = []
        
        for paragraph in self.doc.paragraphs:
            text = paragraph.text.strip()
            
            # Skip empty paragraphs
            if not text:
                continue
                
            # Skip TOC entries (containing tabs and page numbers)
            if '\t' in text and re.search(r'\d+$', text):
                continue
                
            # Check if this is a chapter heading
            is_chapter = False
            chapter_title = None
            
            for pattern, title_func in chapter_patterns:
                match = re.match(pattern, text)
                if match:
                    is_chapter = True
                    if callable(title_func):
                        chapter_title = title_func(match)
                    else:
                        chapter_title = title_func
                    break
            
            if is_chapter and chapter_title:
                # Save previous chapter stats
                if current_chapter:
                    chapters.append(current_chapter)
                    chapter_text = '\n'.join(current_chapter_text)
                    # Count Chinese characters and English words for chapter
                    chinese_chars = len(re.findall(r'[\u4e00-\u9fff]', chapter_text))
                    english_text = re.sub(r'[\u4e00-\u9fff]', ' ', chapter_text)
                    english_words = len([word for word in english_text.split() if word.strip()])
                    word_counts.append(chinese_chars + english_words)
                    equation_counts.append(self._count_equations_in_text(chapter_text))
                    paragraph_counts.append(current_paragraph_count)
                    chapter_paragraphs.append(current_paragraph_count)
                
                # Start new chapter
                current_chapter = chapter_title
                current_chapter_text = []
                current_paragraph_count = 0
            else:
                if current_chapter and text:
                    current_chapter_text.append(text)
                    current_paragraph_count += 1
        
        # Save last chapter
        if current_chapter:
            chapters.append(current_chapter)
            chapter_text = '\n'.join(current_chapter_text)
            # Count Chinese characters and English words for last chapter
            chinese_chars = len(re.findall(r'[\u4e00-\u9fff]', chapter_text))
            english_text = re.sub(r'[\u4e00-\u9fff]', ' ', chapter_text)
            english_words = len([word for word in english_text.split() if word.strip()])
            word_counts.append(chinese_chars + english_words)
            equation_counts.append(self._count_equations_in_text(chapter_text))
            paragraph_counts.append(current_paragraph_count)
            chapter_paragraphs.append(current_paragraph_count)
        
        # Distribute tables and images across chapters based on paragraph counts
        total_chapter_paragraphs = sum(chapter_paragraphs) if chapter_paragraphs else 1
        
        for i, para_count in enumerate(chapter_paragraphs):
            if total_chapter_paragraphs > 0:
                # Distribute proportionally based on paragraph count
                table_portion = round(total_tables * para_count / total_chapter_paragraphs)
                image_portion = round(total_images * para_count / total_chapter_paragraphs)
            else:
                table_portion = 0
                image_portion = 0
            
            table_counts.append(table_portion)
            image_counts.append(image_portion)
        
        return {
            "chapters": chapters,
            "word_counts": word_counts,
            "equation_counts": equation_counts,
            "table_counts": table_counts,
            "image_counts": image_counts,
            "paragraph_counts": paragraph_counts
        }
    
    def _get_latex_content(self) -> str:
        """Convert document to LaTeX using docxlatex"""
        if self.latex_content is None:
            try:
                latex_doc = DocxLatexDocument(str(self.docx_path))
                self.latex_content = latex_doc.get_text()
            except Exception as e:
                print(f"Warning: Failed to convert to LaTeX: {e}")
                self.latex_content = ""
        return self.latex_content
    
    def _count_display_equations(self) -> int:
        """Count display equations converted by docxlatex"""
        latex_content = self._get_latex_content()
        if not latex_content:
            return 0
        
        # Count $$...$$ equations (true display format)
        display_equations = re.findall(r'\$\$[^$]+\$\$', latex_content)
        
        # Count inline equations that are actually display equations
        # These are equations with display math environments like \begin{eqnarray*}, \begin{equation}, etc.
        inline_equations = re.findall(r'\$[^$]+\$', latex_content)
        display_like_patterns = [
            r'\\begin\{eqnarray\*?\}',
            r'\\begin\{equation\*?\}',
            r'\\begin\{align\*?\}',
            r'\\begin\{gather\*?\}',
            r'\\begin\{multline\*?\}',
        ]
        
        display_like_count = 0
        for equation in inline_equations:
            for pattern in display_like_patterns:
                if re.search(pattern, equation):
                    display_like_count += 1
                    break
        
        total_display = len(display_equations) + display_like_count
        print(f"DEBUG: Found {len(display_equations)} true display equations ($$...$$)")
        print(f"DEBUG: Found {display_like_count} display-like equations in inline format")
        print(f"DEBUG: Total display equations: {total_display}")
        
        return total_display
    
    def _count_equations_in_text(self, text: str) -> int:
        """Count display equations in given text using LaTeX conversion"""
        # For chapter-level counting, we need to approximate since we can't easily
        # convert partial text to LaTeX. We'll use a simpler approach for now.
        latex_content = self._get_latex_content()
        if not latex_content:
            return 0
        
        # This is a simplified approach - in a real implementation,
        # you might want to track equation positions in the original document
        # For now, we'll distribute equations proportionally based on text length
        total_text_length = len(self._get_full_text())
        if total_text_length == 0:
            return 0
        
        total_equations = self._count_display_equations()
        text_proportion = len(text) / total_text_length
        
        return round(total_equations * text_proportion)
    
    def _extract_reference_stats(self, text: str) -> Dict[str, Any]:
        """Extract reference statistics"""
        # Find references section - look for the section after 参考文献
        ref_pattern = r'参考文献\s*\n(.+?)(?=\n\s*\n|致谢|附录|\Z)'
        ref_match = re.search(ref_pattern, text, re.DOTALL)
        
        if not ref_match:
            # Try alternative pattern
            ref_pattern = r'参考文献(.+?)(?=\n\s*\n|致谢|附录|\Z)'
            ref_match = re.search(ref_pattern, text, re.DOTALL)
        
        if not ref_match:
            return {
                "total_references": 0,
                "by_indicator": {
                    "期刊论文[J]": 0,
                    "会议论文[C]": 0,
                    "学位论文[D]": 0,
                    "技术报告[R]": 0,
                    "其他": 0
                },
                "by_lang": {"中文文献": 0, "英文文献": 0},
                "recent_3y": 0
            }
        
        ref_text = ref_match.group(1)
        
        # Count total references by looking for numbered reference patterns
        ref_patterns = [
            r'\[\d+\]',  # [1], [2], etc.
            r'^\d+\.',   # 1., 2., etc.
            r'^\d+\s',   # 1 , 2 , etc.
        ]
        
        total_references = 0
        for pattern in ref_patterns:
            matches = re.findall(pattern, ref_text, re.MULTILINE)
            if matches:
                total_references = max(total_references, len(matches))
        
        # Find all reference entries
        ref_lines = [line.strip() for line in ref_text.split('\n') if line.strip()]
        reference_entries = []
        
        # Check for numbered format first
        for line in ref_lines:
            if re.search(r'\[\d+\]', line):  # This is a reference line
                reference_entries.append(line)
        
        # If no numbered format found, treat each non-empty line as a reference
        if not reference_entries:
            reference_entries = [line for line in ref_lines if line.strip()]
            total_references = len(reference_entries)
        
        # Count by indicator
        indicators = {
            "期刊论文[J]": len(re.findall(r'\[J\]', ref_text)),
            "会议论文[C]": len(re.findall(r'\[C\]', ref_text)),
            "学位论文[D]": len(re.findall(r'\[D\]', ref_text)),
            "技术报告[R]": len(re.findall(r'\[R\]', ref_text)),
            "其他": 0
        }
        
        # Calculate "其他"
        total_categorized = sum(indicators.values())
        indicators["其他"] = max(0, total_references - total_categorized)
        
        # Count language for each reference entry based on first word
        chinese_refs = 0
        english_refs = 0
        
        for ref_entry in reference_entries:
            # Remove the reference number [1], [2], etc.
            content = re.sub(r'^\[\d+\]', '', ref_entry).strip()
            
            # Extract the first word (author name)
            first_word_match = re.match(r'^([^\s,，.]+)', content)
            if first_word_match:
                first_word = first_word_match.group(1)
                
                # Check if first word contains Chinese characters
                if re.search(r'[\u4e00-\u9fff]', first_word):
                    chinese_refs += 1
                else:
                    english_refs += 1
            else:
                # If no first word found, assume English
                english_refs += 1
        
        # Count recent 3 years references (2023-2025)
        recent_3y = 0
        for year in range(2023, 2026):  # 2023, 2024, 2025
            recent_3y += len(re.findall(str(year), ref_text))
        
        # Ensure indicators sum equals total_references
        indicator_sum = sum(indicators.values())
        if indicator_sum != total_references:
            diff = total_references - indicator_sum
            indicators["其他"] = max(0, indicators["其他"] + diff)
        
        return {
            "total_references": total_references,
            "by_indicator": indicators,
            "by_lang": {
                "中文文献": chinese_refs,
                "英文文献": english_refs
            },
            "recent_3y": recent_3y
        }


def main():
    import sys
    
    if len(sys.argv) != 2:
        print("Usage: python extract_docx_info.py <docx_file_path>")
        sys.exit(1)
    
    docx_path = sys.argv[1]
    
    try:
        extractor = DocxInfoExtractor(docx_path)
        result = extractor.extract_basic_info()
        print(json.dumps(result, ensure_ascii=False, indent=2))
    except Exception as e:
        print(f"Error processing {docx_path}: {e}")
        sys.exit(1)


if __name__ == "__main__":
    main()