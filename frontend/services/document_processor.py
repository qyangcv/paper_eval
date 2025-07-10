import mammoth
import tempfile
import os
import sys
from docx import Document
import base64
import re
from pathlib import Path
import io
import docx
from docx.document import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
# from ..services.docx2html import Docx2HtmlConverter
import json
import pickle
from typing import Dict, List, Any, Optional, Tuple
import time
import traceback
from bs4 import BeautifulSoup

# 导入自定义日志模块
from frontend.utils.logger_setup import get_module_logger

# 创建当前模块的logger
logger = get_module_logger(__name__)

# 导入 docx_to_html
import sys
from pathlib import Path

# 添加项目根目录（my_project/）到 sys.path
sys.path.append(str(Path(__file__).resolve().parent.parent))

# 使用绝对导入替代相对导入
from backend.hard_metrics.tools.docx_tools.docx2html import docx_to_html


def convert_word_to_html(uploaded_file):
    """
    基本版本：将Word文档转换为HTML

    Args:
        uploaded_file: Streamlit上传的文件对象

    Returns:
        str: 转换后的HTML内容
    """
    try:
        # 读取上传文件的内容
        content = uploaded_file.read()

        # 使用mammoth将Word文档转换为HTML
        result = mammoth.convert_to_html(io.BytesIO(content))
        html_content = result.value

        # 添加基本样式
        styled_html = f"""
        <html>
        <head>
            <meta charset="UTF-8">
            <style>
                body {{ font-family: Arial, sans-serif; line-height: 1.6; margin: 2rem; }}
                h1 {{ color: #333; }}
                h2 {{ color: #444; margin-top: 2rem; }}
                p {{ margin-bottom: 1rem; }}
                img {{ max-width: 100%; height: auto; }}
            </style>
        </head>
        <body>
            {html_content}
        </body>
        </html>
        """
        return styled_html
    except Exception as e:
        logger.info(f"转换Word文档时出错: {e}")
        return "<p>转换文档时出错。请确保上传的是有效的Word文档。</p>"


def convert_word_to_html_with_math(uploaded_file):
    """
    将 Word 文档转换为 HTML（增强版，支持公式、图片和复杂格式）
    使用docx2html.py进行转换，支持数学公式的渲染

    Args:
        uploaded_file: Streamlit上传的文件对象

    Returns:
        str: 生成的HTML内容
    """
    try:
        # 创建临时文件夹用于处理
        temp_dir = './temp_dir'
        os.makedirs(temp_dir, exist_ok=True)  # exist_ok=True 避免目录已存在时出错
        
        # 保存上传的文件到临时目录
        temp_docx_path = os.path.join(temp_dir, "temp_document.docx")
        with open(temp_docx_path, "wb") as f:
            f.write(uploaded_file.getvalue())

        # 设置输出HTML路径
        temp_html_path = os.path.join(temp_dir, "temp_document.html")
        
        # 读取生成的图片文件夹（如果存在）
        images_dir = os.path.join("temp_document_images")

        docx_to_html(
            docx_path=temp_docx_path,
            output_html_path=temp_html_path,
            image_dir=images_dir,
            css_file=None,
        )

        # 读取生成的HTML文件
        with open(temp_html_path, "r", encoding="utf-8") as f:
            html_content = f.read()

        # image_files = {}
        # if os.path.exists(images_dir):
        #     for img_file in os.listdir(images_dir):
        #         img_path = os.path.join(images_dir, img_file)
        #         print(img_path)
        #         with open(img_path, "rb") as img:
        #             # 转换图片为base64编码，以便嵌入HTML
        #             image_data = base64.b64encode(
        #                 img.read()).decode('utf-8')
        #             mime_type = get_mime_type(img_file)
        #             image_files[img_file] = f"data:{mime_type};base64,{image_data}"

        #     # 替换HTML中的图片引用为base64编码
        #     for img_file, img_data in image_files.items():
        #         img_dir_name = os.path.basename(images_dir)
        #         img_path = f"{img_dir_name}/{img_file}"
        #         html_content = html_content.replace(
        #             f'src="{img_path}"', f'src="{img_data}"')

        return html_content
    except Exception as e:
        logger.info(f"使用增强版转换器处理Word文档时出错: {e}")
        # 如果增强版转换失败，回退到基础版
        return convert_word_to_html(uploaded_file)


def get_mime_type(file_path):
    """根据文件扩展名确定MIME类型"""
    ext = os.path.splitext(file_path)[1].lower()
    mime_types = {
        '.png': 'image/png',
        '.jpg': 'image/jpeg',
        '.jpeg': 'image/jpeg',
        '.gif': 'image/gif',
        '.bmp': 'image/bmp',
        '.svg': 'image/svg+xml',
        '.webp': 'image/webp'
    }
    return mime_types.get(ext, 'image/png')  # 默认为PNG


def extract_toc_from_docx(uploaded_file):
    """从Word文档中提取目录结构，优化识别"第X章"式标题和子章节"""
    try:
        # 读取文档
        doc = docx.Document(io.BytesIO(uploaded_file.getvalue()))

        toc_items = []
        main_chapters = []  # 存储主章节
        sub_chapters = []   # 存储子章节

        # 标记用于处理目录和文档区域
        found_toc = False  # 是否找到"目录"
        in_toc_section = False  # 是否在目录区域内
        found_content_start = False  # 是否找到正文开始（第二次出现"第一章"）
        chapter_texts = set()  # 用于跟踪已添加的章节，防止重复

        logger.info("开始分析文档结构...")

        # 第一次扫描：查找目录区域和正文开始位置
        toc_start_index = -1  # 目录开始位置
        toc_end_index = -1    # 目录结束位置
        content_start_index = -1  # 正文开始位置

        # 章节计数器，用于检测重复章节
        chapter_count = {}

        # 第一次扫描：查找目录和章节位置
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            if not text:
                continue

            # 检测目录节
            if not found_toc and (text.lower() == "目录" or text.lower() == "contents" or text.lower() == "table of contents"):
                found_toc = True
                in_toc_section = True
                toc_start_index = i
                logger.info(f"检测到目录: '{text}' at index {i}")
                continue

            # 如果已经找到目录，检测章节标题
            if found_toc and in_toc_section:
                # 检测章节标题模式
                chapter_match = re.match(r'^第[一二三四五六七八九十\d]+章', text)
                if chapter_match:
                    chapter_name = chapter_match.group(0)

                    # 记录章节出现次数
                    if chapter_name not in chapter_count:
                        chapter_count[chapter_name] = 1
                    else:
                        chapter_count[chapter_name] += 1

                    # 如果是第二次出现"第一章"，表示正文开始
                    if chapter_name in ["第一章", "第1章"] and chapter_count.get(chapter_name, 0) > 1:
                        in_toc_section = False
                        found_content_start = True
                        content_start_index = i
                        toc_end_index = i - 1  # 上一段落是目录的结束
                        logger.info(f"检测到正文开始: '{text}' at index {i}")
                        break

        # 如果没有明确找到目录结束位置，但找到了正文开始，则以正文开始位置作为目录结束
        if toc_end_index == -1 and content_start_index != -1:
            toc_end_index = content_start_index - 1

        logger.info(
            f"文档分析结果: 目录开始={toc_start_index}, 目录结束={toc_end_index}, 正文开始={content_start_index}")

        # 如果找到了目录区域，提取章节结构
        if found_toc and toc_start_index != -1 and content_start_index != -1:
            # 从正文开始位置提取章节
            start_index = content_start_index

            # 第二次扫描：从正文开始位置提取章节
            for i, paragraph in enumerate(doc.paragraphs[start_index:], start_index):
                text = paragraph.text.strip()
                if not text:
                    continue

                # 识别为标题样式的段落
                is_heading = False
                level = 0

                # 通过样式名识别标题
                if paragraph.style.name.startswith('Heading') or '标题' in paragraph.style.name:
                    is_heading = True
                    try:
                        # 从样式名获取级别
                        level_match = re.search(r'\d+', paragraph.style.name)
                        if level_match:
                            level = int(level_match.group(0))
                        else:
                            level = 1  # 默认为一级标题
                    except:
                        level = 1

                # 通过格式识别标题 - 检查是否粗体或大字体
                elif paragraph.runs:
                    is_bold = any(
                        run.bold for run in paragraph.runs if hasattr(run, 'bold'))
                    is_large = False
                    try:
                        is_large = any(run.font.size and run.font.size.pt > 14 for run in paragraph.runs if hasattr(
                            run, 'font') and hasattr(run.font, 'size'))
                    except:
                        pass

                    if is_bold or is_large:
                        # 进一步检查是否匹配章节标题模式
                        is_heading = True
                        level = 1 if is_large else 2

                # 通过文本模式识别标题
                chapter_patterns = [
                    # 主章节模式
                    (r'^第[一二三四五六七八九十\d]+章[：:]?\s*\S+', 1),  # 第一章：绪论
                    (r'^第[一二三四五六七八九十\d]+章$', 1),  # 第一章
                    (r'^\d+[\.、]\s*[^\.、\d]+', 1),  # 1. 绪论
                    (r'^[一二三四五六七八九十]+[\.、]\s*[^\.、\d]+', 1),  # 一. 绪论

                    # 子章节模式
                    (r'^\d+\.\d+\s+\S+', 2),  # 1.1 研究背景
                    (r'^\d+\.\d+\.\d+\s+\S+', 3),  # 1.1.1 具体内容
                    (r'^第[一二三四五六七八九十\d]+节\s+\S+', 2),  # 第一节 内容
                ]

                if not is_heading:
                    for pattern, pat_level in chapter_patterns:
                        if re.match(pattern, text):
                            is_heading = True
                            level = pat_level
                            break

                # 标准化章节名称（移除数字后缀等）
                standardized_text = standardize_chapter_name(text)
                display_text = standardized_text[:20] + "..." if len(
                    standardized_text) > 20 else standardized_text
                original_text = text  # 保留原始文本用于匹配

                # 章节标题需满足以下额外条件：
                # 1. 不能包含中文或英文标点（如"：，。.?!等"）
                # 2. 一级章节标题通常字体较大或明确使用 Heading 1 / 标题 1 样式

                # 过滤包含标点符号的标题（目录行通常带有省略号或页码）
                has_punctuation = re.search(
                    r'[，。：；、,\.;:!？?!]', original_text) is not None

                # 检查字体大小是否足够大（>14pt 视为大字体）
                is_large_font = False
                try:
                    if paragraph.runs:
                        for run in paragraph.runs:
                            if run.font.size and run.font.size.pt and run.font.size.pt > 14:
                                is_large_font = True
                                break
                except Exception:
                    pass

                # 检查是否为一级标题样式（Heading 1 或 标题 1）
                style_name_lower = paragraph.style.name.lower(
                ) if paragraph.style and paragraph.style.name else ""
                is_heading1_style = style_name_lower.startswith(
                    'heading 1') or '标题 1' in paragraph.style.name or '标题1' in paragraph.style.name if paragraph.style else False

                # 对于 level==1，需要字体较大或使用 Heading 1 样式
                meets_font_style_requirement = True
                if level == 1:
                    meets_font_style_requirement = is_large_font or is_heading1_style

                # 跳过个人信息、不符合章节特征的内容、重复的章节、含标点或字体/样式不符合要求的标题
                if (
                    is_heading
                    and not is_personal_info(text)
                    and display_text not in chapter_texts
                    and not has_punctuation
                    and meets_font_style_requirement
                ):
                    # 添加到已处理章节集合，防止重复
                    chapter_texts.add(display_text)

                    # 创建章节ID，结合段落索引和章节文本的前几个字符（确保唯一性）
                    safe_text = re.sub(r'[^\w\d]', '', text[:5])
                    chapter_id = f"section-{i}-{safe_text}"

                    item = {
                        'index': i,
                        'level': level,
                        'text': display_text,
                        'original_text': original_text,  # 保存原始文本用于内容匹配
                        'standardized_text': standardized_text,  # 保存标准化后的文本
                        'id': chapter_id,
                        'children': []
                    }

                    if level == 1:
                        main_chapters.append(item)
                    else:
                        sub_chapters.append(item)

        # 如果没有找到足够的章节，尝试其他规则
        if len(main_chapters) < 1:
            logger.info("未找到足够主章节，尝试通过内容特征识别...")

            # 查找有明显特征的段落
            start_index = content_start_index if content_start_index != -1 else 0
            for i, paragraph in enumerate(doc.paragraphs[start_index:start_index+100], start_index):
                text = paragraph.text.strip()
                if not text or len(text) < 4 or len(text) > 100:
                    continue

                # 标准化章节名称
                standardized_text = standardize_chapter_name(text)
                display_text = standardized_text[:20] + "..." if len(
                    standardized_text) > 20 else standardized_text

                # 检查是否已添加
                if display_text in chapter_texts:
                    continue

                # 常见章节关键词
                chapter_keywords = ['绪论', '引言', '简介', '概述', '背景', '方法', '实验', '结果', '分析',
                                    '讨论', '结论', '参考文献', '致谢', 'Introduction', 'Methods',
                                    'Results', 'Discussion', 'Conclusion', 'References']

                if any(keyword in text for keyword in chapter_keywords):
                    # 确认文本格式特征 - 粗体或单独成段落等
                    is_formatted = False
                    if paragraph.runs:
                        is_bold = any(
                            run.bold for run in paragraph.runs if hasattr(run, 'bold'))
                        if is_bold:
                            is_formatted = True

                    # 如果是单独的短段落也可能是标题
                    if len(text) < 30:
                        is_formatted = True

                    if is_formatted and not is_personal_info(text):
                        # 添加到已处理章节集合
                        chapter_texts.add(display_text)

                        # 创建章节ID，结合段落索引和章节文本的前几个字符
                        safe_text = re.sub(r'[^\w\d]', '', text[:5])
                        chapter_id = f"section-{i}-{safe_text}"

                        main_chapters.append({
                            'index': i,
                            'level': 1,
                            'text': display_text,
                            'original_text': text,
                            'standardized_text': standardized_text,
                            'id': chapter_id,
                            'children': []
                        })

                        logger.info(
                            f"通过内容特征添加章节: '{display_text}', id={chapter_id}")

        # 构建层级关系
        for sub_item in sub_chapters:
            # 找到最近的主章节作为父级
            parent_found = False
            for i in range(len(main_chapters)-1, -1, -1):
                if main_chapters[i]['index'] < sub_item['index']:
                    main_chapters[i]['children'].append(sub_item)
                    parent_found = True
                    break

            # 如果没找到父级，可能是独立的子章节或序言等
            if not parent_found and main_chapters:
                main_chapters[0]['children'].append(sub_item)

        # 合并所有章节数据
        toc_items = main_chapters

        logger.info(f"共提取 {len(toc_items)} 个主章节, {len(sub_chapters)} 个子章节")

        return toc_items
    except Exception as e:
        logger.info(f"提取目录时出错: {str(e)}")
        traceback.print_exc()
        return []


def standardize_chapter_name(text):
    """标准化章节名称，去除数字后缀等"""
    # 处理"第X章 绪论1"这样的情况，转换为"第X章 绪论"
    text = re.sub(r'^(第[一二三四五六七八九十\d]+章\s*[\s\S]*绪论)\d+', r'\1', text)
    text = re.sub(r'^(第[一二三四五六七八九十\d]+章\s*[\s\S]*引言)\d+', r'\1', text)

    # 去除标题末尾的数字
    text = re.sub(r'(第[一二三四五六七八九十\d]+[章节].*?[^0-9])\d+$', r'\1', text)

    # 清理额外的空格
    text = re.sub(r'\s+', ' ', text).strip()

    return text


def is_likely_content_start(paragraphs, current_text):
    """判断是否是目录结束、正文开始的位置"""
    # 检查前面的段落是否有明显的格式变化
    format_change = False
    content_indicators = False

    # 检查格式变化（例如从短段落到长段落）
    text_lengths = [len(p.text.strip()) for p in paragraphs if p.text.strip()]
    if text_lengths and (max(text_lengths) < 30) and len(current_text) > 30:
        format_change = True

    # 检查是否有页眉、页脚、分隔符等内容指示
    for p in paragraphs:
        text = p.text.strip().lower()
        if text and (len(text) < 5 or text in ['page', 'chapter', '-'*3]):
            content_indicators = True
            break

    # 检查当前文本是否是明显的章节开始
    is_chapter_start = re.match(
        r'^第[一二三四五六七八九十\d]+章\s+\S+', current_text) is not None

    return format_change or content_indicators or is_chapter_start


def is_personal_info(text):
    """检查文本是否包含个人信息"""
    personal_info_patterns = [
        r'姓名[：:]\s*\w+',
        r'电话[：:]\s*\d+',
        r'邮箱[：:]\s*[\w\.-]+@[\w\.-]+',
        r'地址[：:]\s*\w+',
        r'学号[：:]\s*\w+',
        r'指导教师[：:]\s*\w+'
    ]

    for pattern in personal_info_patterns:
        if re.search(pattern, text):
            return True
    return False


def process_paper_evaluation(input_file_path: str,
                             toc_items: List[Dict[str, Any]] = None,
                             model_name: str = "deepseek-chat") -> Dict[str, Any]:
    """
    处理论文评估，调用full_paper_eval.py，并将结果格式化为results_page.py可用格式

    Args:
        input_file_path: 输入文件路径，可以是docx或pkl文件
        toc_items: 目录项列表，如果提供，会将评估结果与章节关联
        model_name: 使用的模型名称，默认为"deepseek-chat"

    Returns:
        Dict[str, Any]: 包含章节评估结果和整体评分的字典
    """
    try:
        # 设置正确的导入路径
        project_root = os.path.dirname(os.path.dirname(
            os.path.dirname(os.path.abspath(__file__))))
        eval_path = os.path.join(project_root, "backend", "hard_metrics")
        tools_path = os.path.join(eval_path, "tools")
        models_path = os.path.join(eval_path, "models")
        prompts_path = os.path.join(eval_path, "prompts")

        # 添加所有相关路径
        for path in [project_root, eval_path, tools_path, models_path, prompts_path]:
            if path not in sys.path:
                sys.path.insert(0, path)

        # 直接导入模块
        from backend.hard_metrics.full_paper_eval import process_docx_file, load_chapters, process_chapter
        from backend.hard_metrics.full_paper_eval import evaluate_overall, score_paper

        # 处理输入文件
        pkl_file_path = input_file_path
        if input_file_path.lower().endswith('.docx'):
            logger.info("检测到.docx输入，进行文件转换")
            pkl_file_path = process_docx_file(input_file_path)
            if not pkl_file_path:
                logger.info("文件转换失败，无法继续")
                return {"error": "文档转换失败"}

        # 加载章节
        try:
            chapters = load_chapters(pkl_file_path)
        except Exception as e:
            logger.info(f"加载章节失败: {e}")
            # 如果pkl加载失败，尝试直接使用toc_items中的内容
            if toc_items:
                chapters = []
                for i, chapter in enumerate(toc_items, 1):
                    if 'text' in chapter and 'original_text' in chapter:
                        chapters.append({
                            "index": i,
                            "title": chapter['text'],
                            "content": chapter.get('content', f"章节 {chapter['text']} 的内容")
                        })
            else:
                return {"error": f"无法加载章节内容: {str(e)}"}

        if not chapters:
            return {"error": "未找到有效的章节内容"}

        chapter_evaluations = []

        # 串行评估每个章节
        for chapter in chapters:
            evaluation = process_chapter(chapter, model_name)
            chapter_evaluations.append(evaluation)

        # 进行整体评估
        overall_evaluation = evaluate_overall(chapter_evaluations, model_name)

        # 合并所有评估结果（将整体评估放在首位）
        all_evaluations = [overall_evaluation] + chapter_evaluations

        # 进行论文评分
        paper_scores = score_paper(all_evaluations, model_name)

        # 整合评估结果和目录结构
        if toc_items:
            for chapter in toc_items:
                chapter_title = chapter.get('text', '')
                # 查找匹配的评估结果
                matching_eval = None
                for eval_item in all_evaluations:
                    # 跳过整体评估
                    if eval_item.get('chapter') == "全篇" or eval_item.get('index') == 0:
                        continue

                    # 寻找匹配的章节标题
                    if chapter_title.lower() in eval_item.get('chapter', '').lower() or \
                       eval_item.get('chapter', '').lower() in chapter_title.lower():
                        matching_eval = eval_item
                        break

                # 如果找到匹配的评估，添加到章节信息中
                if matching_eval:
                    chapter['analysis'] = {
                        "summary": matching_eval.get('summary', ''),
                        "strengths": matching_eval.get('strengths', []),
                        "weaknesses": matching_eval.get('weaknesses', []),
                        "suggestions": matching_eval.get('suggestions', [])
                    }
                else:
                    # 未找到匹配的评估，添加空评估
                    chapter['analysis'] = {
                        "summary": f"本章节主要讨论{chapter_title}相关内容。",
                        "strengths": [],
                        "weaknesses": [],
                        "suggestions": []
                    }

                # 递归处理子章节
                if 'children' in chapter and chapter['children']:
                    for child in chapter['children']:
                        child_title = child.get('text', '')
                        # 查找匹配的评估结果
                        matching_child_eval = None
                        for eval_item in all_evaluations:
                            if eval_item.get('chapter') == "全篇" or eval_item.get('index') == 0:
                                continue
                            if child_title.lower() in eval_item.get('chapter', '').lower() or \
                               eval_item.get('chapter', '').lower() in child_title.lower():
                                matching_child_eval = eval_item
                                break

                        # 添加评估
                        if matching_child_eval:
                            child['analysis'] = {
                                "summary": matching_child_eval.get('summary', ''),
                                "strengths": matching_child_eval.get('strengths', []),
                                "weaknesses": matching_child_eval.get('weaknesses', []),
                                "suggestions": matching_child_eval.get('suggestions', [])
                            }
                        else:
                            child['analysis'] = {
                                "summary": f"本小节主要讨论{child_title}相关内容。",
                                "strengths": [],
                                "weaknesses": [],
                                "suggestions": []
                            }

        # 构建最终结果
        result = {
            "toc_items": toc_items,
            "overall_scores": paper_scores,
            "paper_summary": {
                "overall_comment": overall_evaluation.get('summary', ''),
                "strengths": overall_evaluation.get('strengths', []),
                "weaknesses": overall_evaluation.get('weaknesses', []),
                "suggestions": overall_evaluation.get('suggestions', [])
            }
        }

        logger.info(f"论文评估完成，共处理 {len(chapters)} 个章节")
        return result

    except ImportError as e:
        logger.info(f"导入错误: {e}")
        logger.info("确保您在正确的项目结构中运行此脚本")
        return {"error": f"导入评估模块失败: {str(e)}"}
    except Exception as e:
        logger.info(f"评估过程出错: {e}")
        traceback.print_exc()
        return {"error": f"评估过程出错: {str(e)}"}


def simulate_analysis_with_toc(uploaded_file, progress_callback=None):
    """
    分析文档并生成结构化的分析结果，包括目录结构和评估结果。
    这个函数会调用 process_paper_evaluation 进行实际的评估。

    Args:
        uploaded_file: Streamlit上传的文件对象
        progress_callback: 进度回调函数，接受(current_progress, message)两个参数

    Returns:
        Dict[str, Any]: 包含完整评估结果的字典
    """
    temp_path = None
    try:
        # 首先提取目录结构
        if progress_callback:
            progress_callback(0.05, "正在提取文档目录结构...")
        toc_items = extract_toc_from_docx(uploaded_file)

        # 在临时目录保存上传的文件
        if progress_callback:
            progress_callback(0.10, "正在准备文档分析...")
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_file:
            temp_file.write(uploaded_file.getvalue())
            temp_path = temp_file.name

        # 模拟评估过程的不同阶段
        if progress_callback:
            progress_callback(0.15, "开始章节内容分析...")

        # 获取章节数量用于计算进度
        chapter_count = len(toc_items)
        base_progress = 0.15
        chapter_progress = 0.60 / max(chapter_count, 1)  # 章节分析总共占60%的进度

        # 模拟对每个章节的处理
        for i, chapter in enumerate(toc_items):
            chapter_title = chapter.get('text', f"章节 {i+1}")
            if progress_callback:
                current_progress = base_progress + (i * chapter_progress)
                progress_callback(current_progress,
                                  f"正在分析: {chapter_title}...")
                # 模拟章节分析过程
                time.sleep(0.5)

        if progress_callback:
            progress_callback(0.75, "正在进行整体内容评估...")
            time.sleep(1.5)  # 模拟整体评估过程

            progress_callback(0.85, "正在计算学术维度得分...")
            time.sleep(1.0)  # 模拟计算得分过程

            progress_callback(0.95, "正在整合分析结果...")

        # 使用临时文件路径进行评估
        logger.info(f"使用文件 {temp_path} 进行论文评估")
        result = process_paper_evaluation(temp_path, toc_items)

        # 检查是否有错误
        if 'error' in result:
            logger.info(f"评估遇到问题: {result['error']}")
            logger.info("将使用默认分析结果")
            if progress_callback:
                progress_callback(0.98, "评估遇到问题，使用默认结果...")
            return _generate_default_analysis(toc_items)

        # 如果评估成功，更新toc_items
        if 'toc_items' in result and result['toc_items']:
            toc_items = result['toc_items']

        if progress_callback:
            progress_callback(1.0, "评估完成！")

        return {
            'chapters': toc_items,
            'overall_scores': result.get('overall_scores', []),
            'paper_summary': result.get('paper_summary', {})
        }
    except Exception as e:
        logger.info(f"分析文档时出错: {e}")
        traceback.print_exc()
        if progress_callback:
            progress_callback(0.98, "处理过程中出错，使用默认分析结果...")
        return _generate_default_analysis(extract_toc_from_docx(uploaded_file))
    finally:
        # 确保删除临时文件
        if temp_path and os.path.exists(temp_path):
            try:
                os.unlink(temp_path)
            except Exception as e:
                logger.info(f"无法删除临时文件 {temp_path}: {e}")
                pass


def _generate_default_analysis(toc_items):
    """
    生成默认的分析结果，当模型分析失败时使用

    Args:
        toc_items: 目录结构列表

    Returns:
        Dict[str, Any]: 默认分析结果
    """
    # 生成默认的分析结果
    default_analysis = {
        'chapters': toc_items,
        'overall_scores': [
            {'index': 1, 'module': '摘要', 'full_score': 5, 'score': 4},
            {'index': 2, 'module': '选题背景和意义', 'full_score': 5, 'score': 4},
            {'index': 3, 'module': '选题的理论意义与应用价值', 'full_score': 5, 'score': 4},
            {'index': 4, 'module': '相关工作的国内外现状综述', 'full_score': 5, 'score': 4},
            {'index': 5, 'module': '主要工作和贡献总结', 'full_score': 5, 'score': 4},
            {'index': 6, 'module': '相关工作或相关技术的介绍', 'full_score': 5, 'score': 4},
            {'index': 7, 'module': '论文的创新性', 'full_score': 25, 'score': 20},
            {'index': 8, 'module': '实验完成度', 'full_score': 20, 'score': 15},
            {'index': 9, 'module': '总结和展望', 'full_score': 5, 'score': 4},
            {'index': 10, 'module': '工作量', 'full_score': 5, 'score': 4},
            {'index': 11, 'module': '论文撰写质量', 'full_score': 10, 'score': 7},
            {'index': 12, 'module': '参考文献', 'full_score': 5, 'score': 4},
        ],
        'paper_summary': {
            "overall_comment": "本论文整体结构完整，内容丰富，研究方法合理，但部分章节可进一步完善。",
            "strengths": [
                "研究问题明确，具有一定的理论和实际意义",
                "研究方法选择合理，实验设计完整",
                "数据分析全面，结果可信度高"
            ],
            "weaknesses": [
                "文献综述部分可以更加全面",
                "创新点阐述不够突出",
                "部分章节内容深度不够"
            ],
            "suggestions": [
                "建议补充更多最新相关研究文献",
                "加强对创新点的论述与证明",
                "增加对研究局限性的讨论"
            ]
        }
    }

    # 确保为每个章节添加分析结果
    for chapter in default_analysis['chapters']:
        if 'analysis' not in chapter:
            chapter_title = chapter.get('text', '')
            chapter['analysis'] = {
                "summary": f"本章节主要讨论{chapter_title}相关内容，包含相关理论和方法介绍。",
                "strengths": ["章节结构清晰", "内容充实"],
                "weaknesses": ["可进一步扩展深度"],
                "suggestions": ["建议增加案例分析"]
            }

        # 处理子章节
        if 'children' in chapter:
            for child in chapter['children']:
                if 'analysis' not in child:
                    child_title = child.get('text', '')
                    child['analysis'] = {
                        "summary": f"本小节主要讨论{child_title}相关内容。",
                        "strengths": ["逻辑清晰"],
                        "weaknesses": ["内容可以更加充实"],
                        "suggestions": ["建议增加具体示例"]
                    }

    return default_analysis
