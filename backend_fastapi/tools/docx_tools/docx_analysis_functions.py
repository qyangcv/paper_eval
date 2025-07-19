#!/usr/bin/env python3
"""
Four separate functions extracted from extract_docx_info.py for different analysis tasks.
"""

import json
import re
from typing import Dict, Any

try:
    from docx import Document
    from docx.text.paragraph import Paragraph
    from docx.table import Table
    from docx.oxml.text.paragraph import CT_P
    from docx.oxml.table import CT_Tbl
    from docxlatex.docxlatex import Document as DocxLatexDocument
except ImportError as e:
    print(f"Please install required packages: {e}")
    print("Run: pip install python-docx docxlatex")
    exit(1)


def _get_full_text(doc) -> str:
    """Get full text content from document"""
    full_text = []
    for element in doc.element.body:
        if isinstance(element, CT_P):
            paragraph = Paragraph(element, doc)
            full_text.append(paragraph.text)
        elif isinstance(element, CT_Tbl):
            table = Table(element, doc)
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
    return '\n'.join(full_text)


def _get_latex_content(docx_path: str) -> str:
    """Convert document to LaTeX using docxlatex"""
    try:
        latex_doc = DocxLatexDocument(str(docx_path))
        return latex_doc.get_text()
    except Exception as e:
        print(f"Warning: Failed to convert to LaTeX: {e}")
        return ""


def _count_display_equations(docx_path: str) -> int:
    """Count display equations converted by docxlatex"""
    latex_content = _get_latex_content(docx_path)
    if not latex_content:
        return 0
    
    # Count $$...$$ equations (true display format)
    display_equations = re.findall(r'\$\$[^$]+\$\$', latex_content)
    
    # Count inline equations that are actually display equations
    # These are equations with display math environments like \begin{eqnarray*}, \begin{equation}, etc.
    inline_equations = re.findall(r'\$[^$]+\$', latex_content)
    display_like_patterns = [
        r'\\begin\{eqnarray\*?\}',
        r'\\begin\{equation\*?\}',
        r'\\begin\{align\*?\}',
        r'\\begin\{gather\*?\}',
        r'\\begin\{multline\*?\}',
    ]
    
    display_like_count = 0
    for equation in inline_equations:
        for pattern in display_like_patterns:
            if re.search(pattern, equation):
                display_like_count += 1
                break
    
    total_display = len(display_equations) + display_like_count
    return total_display


def _count_equations_in_text(docx_path: str, text: str, full_text: str) -> int:
    """Count display equations in given text using improved estimation"""
    # First try to count equations directly in the text
    # Look for equation-like patterns in the text
    equation_patterns = [
        r'\$[^$]+\$',  # Inline equations
        r'\$\$[^$]+\$\$',  # Display equations
        r'\\begin\{equation\}.*?\\end\{equation\}',  # LaTeX equation environment
        r'\\begin\{align\}.*?\\end\{align\}',  # LaTeX align environment
        r'\\[.*?\\]',  # LaTeX display math
        r'\([0-9]+\.[0-9]+\)',  # Numbered equations like (1.1)
    ]
    
    direct_count = 0
    for pattern in equation_patterns:
        matches = re.findall(pattern, text, re.DOTALL)
        direct_count += len(matches)
    
    # If direct counting finds equations, use that
    if direct_count > 0:
        return direct_count
    
    # Otherwise, use the original proportional method as fallback
    latex_content = _get_latex_content(docx_path)
    if not latex_content:
        return 0
    
    total_text_length = len(full_text)
    if total_text_length == 0:
        return 0
    
    total_equations = _count_display_equations(docx_path)
    text_proportion = len(text) / total_text_length
    
    return max(0, round(total_equations * text_proportion))


def get_basic_info(docx_path: str) -> Dict[str, Any]:
    """Extract basic information from the docx file"""
    doc = Document(docx_path)
    text_content = _get_full_text(doc)
    
    # Extract title
    title_pattern = r'题目[：:]?\s*([^\n]+)'
    title_match = re.search(title_pattern, text_content)
    title = title_match.group(1).strip() if title_match else "未识别"
    
    # Extract author
    author_patterns = [
        r'学\s*生[：:]?\s*([^\s\n：:]+)',
        r'作\s*者[：:]?\s*([^\s\n：:]+)',
        r'姓\s*名[：:]?\s*([^\s\n：:]+)',
    ]
    
    author = "未识别"
    for pattern in author_patterns:
        match = re.search(pattern, text_content)
        if match:
            author_candidate = match.group(1).strip()
            # Filter out long text that's clearly not a name
            if len(author_candidate) < 20 and not re.search(r'[。，！？]', author_candidate):
                author = author_candidate
                break
    
    # Extract school
    school_patterns = [
        r'学\s*院[：:]?\s*([^\s\n]+)',
        r'院\s*系[：:]?\s*([^\s\n]+)',
        r'系\s*别[：:]?\s*([^\s\n]+)',
    ]
    
    school = "未识别"
    for pattern in school_patterns:
        match = re.search(pattern, text_content)
        if match:
            school = match.group(1).strip()
            break
    
    # Extract advisor
    advisor_patterns = [
        r'指导教师[：:]?\s*([^\s\n]+)',
        r'导\s*师[：:]?\s*([^\s\n]+)',
        r'指导老师[：:]?\s*([^\s\n]+)',
    ]
    
    advisor = "未识别"
    for pattern in advisor_patterns:
        match = re.search(pattern, text_content)
        if match:
            advisor = match.group(1).strip()
            break
    
    # Extract keywords
    keywords_pattern = r'关键词[：:]?\s*([^\n]+)'
    keywords_match = re.search(keywords_pattern, text_content)
    keywords = []
    if keywords_match:
        keywords_text = keywords_match.group(1).strip()
        # Split by common delimiters
        keywords = re.split(r'[；;，,、\s]+', keywords_text)
        keywords = [kw.strip() for kw in keywords if kw.strip()]
    
    return {
        "title": title,
        "author": author,
        "school": school,
        "advisor": advisor,
        "keywords": keywords
    }


def get_overall_stats(docx_path: str) -> Dict[str, int]:
    """Extract overall statistics"""
    doc = Document(docx_path)
    
    total_words = 0
    total_equations = 0
    total_paragraphs = 0
    total_images = 0
    total_tables = 0
    
    # Count paragraphs and words (Chinese + English)
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            total_paragraphs += 1
            # Count Chinese characters and English words
            text = paragraph.text
            # Count Chinese characters
            chinese_chars = len(re.findall(r'[\u4e00-\u9fff]', text))
            # Count English words (split by whitespace, filter out Chinese)
            english_text = re.sub(r'[\u4e00-\u9fff]', ' ', text)
            english_words = len([word for word in english_text.split() if word.strip()])
            total_words += chinese_chars + english_words
    
    # Count tables
    total_tables = len(doc.tables)
    
    # Count images (inline shapes)
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run.element.xpath('.//a:blip'):
                total_images += 1
    
    # Count display equations using LaTeX conversion
    total_equations = _count_display_equations(docx_path)
    
    return {
        "total_words": total_words,
        "total_equations": total_equations,
        "total_paragraphs": total_paragraphs,
        "total_images": total_images,
        "total_tables": total_tables
    }


def get_chapter_stats(docx_path: str) -> Dict[str, Any]:
    """Extract chapter-based statistics"""
    doc = Document(docx_path)
    full_text = _get_full_text(doc)
    
    chapters = []
    word_counts = []
    equation_counts = []
    table_counts = []
    image_counts = []
    paragraph_counts = []
    
    # Common chapter patterns - support both Chinese and Arabic numbers
    chapter_patterns = [
        (r'^(摘\s*要)$', '摘要'),
        (r'^(ABSTRACT)$', 'Abstract'),
        # Chinese numbers
        (r'^第(一|二|三|四|五|六|七|八|九|十)章\s+(.+)', lambda m: f'第{m.group(1)}章 {m.group(2)}'),
        (r'^第(一|二|三|四|五|六|七|八|九|十)章$', lambda m: f'第{m.group(1)}章'),
        # Arabic numbers
        (r'^第(\d+)章\s+(.+)', lambda m: f'第{m.group(1)}章 {m.group(2)}'),
        (r'^第(\d+)章$', lambda m: f'第{m.group(1)}章'),
        # Simple number patterns (for titles like "1 绪论", "1.1 研究背景")
        (r'^(\d+)\s+([^.\d].+)', lambda m: f'第{m.group(1)}章 {m.group(2)}'),
        (r'^(\d+)\.(\d+)?\s+(.+)', lambda m: f'第{m.group(1)}章 {m.group(3)}'),
        (r'^(参考文献)$', '参考文献'),
    ]
    
    # Get total counts for distribution
    total_tables = len(doc.tables)
    total_images = 0
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run.element.xpath('.//a:blip'):
                total_images += 1
    
    current_chapter = None
    current_chapter_text = []
    current_paragraph_count = 0
    chapter_paragraphs = []
    
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        
        # Skip empty paragraphs
        if not text:
            continue
            
        # Skip TOC entries (more precise filtering)
        # TOC entries usually have: title + tabs/dots + page number
        if (('\t' in text and re.search(r'\d+$', text)) or 
            (re.search(r'\.{3,}', text) and re.search(r'\d+$', text)) or
            (len(text) < 20 and re.search(r'^\d+\s*$', text))):  # Pure page numbers
            continue
            
        # Check if this is a chapter heading
        is_chapter = False
        chapter_title = None
        
        for pattern, title_func in chapter_patterns:
            match = re.match(pattern, text)
            if match:
                is_chapter = True
                if callable(title_func):
                    chapter_title = title_func(match)
                else:
                    chapter_title = title_func
                break
        
        if is_chapter and chapter_title:
            # Save previous chapter stats
            if current_chapter:
                chapters.append(current_chapter)
                chapter_text = '\n'.join(current_chapter_text)
                # Count Chinese characters and English words for chapter
                chinese_chars = len(re.findall(r'[\u4e00-\u9fff]', chapter_text))
                english_text = re.sub(r'[\u4e00-\u9fff]', ' ', chapter_text)
                english_words = len([word for word in english_text.split() if word.strip()])
                word_counts.append(chinese_chars + english_words)
                equation_counts.append(_count_equations_in_text(docx_path, chapter_text, full_text))
                paragraph_counts.append(current_paragraph_count)
                chapter_paragraphs.append(current_paragraph_count)
            
            # Start new chapter
            current_chapter = chapter_title
            current_chapter_text = []
            current_paragraph_count = 0
        else:
            if current_chapter and text:
                current_chapter_text.append(text)
                current_paragraph_count += 1
    
    # Save last chapter
    if current_chapter:
        chapters.append(current_chapter)
        chapter_text = '\n'.join(current_chapter_text)
        # Count Chinese characters and English words for last chapter
        chinese_chars = len(re.findall(r'[\u4e00-\u9fff]', chapter_text))
        english_text = re.sub(r'[\u4e00-\u9fff]', ' ', chapter_text)
        english_words = len([word for word in english_text.split() if word.strip()])
        word_counts.append(chinese_chars + english_words)
        equation_counts.append(_count_equations_in_text(docx_path, chapter_text, full_text))
        paragraph_counts.append(current_paragraph_count)
        chapter_paragraphs.append(current_paragraph_count)
    
    # Distribute tables and images across chapters more intelligently
    # Try to count actual tables and images per chapter first
    chapter_table_counts = []
    chapter_image_counts = []
    
    # Initialize with zeros
    for _ in chapters:
        chapter_table_counts.append(0)
        chapter_image_counts.append(0)
    
    # Try to assign tables to chapters based on proximity
    # This is a simplified approach - ideally we'd track exact positions
    if chapters:
        tables_per_chapter = max(1, total_tables // len(chapters))
        images_per_chapter = max(1, total_images // len(chapters))
        
        # Distribute evenly with remainder going to first chapters
        for i in range(len(chapters)):
            if i < total_tables % len(chapters):
                chapter_table_counts[i] = tables_per_chapter + 1
            else:
                chapter_table_counts[i] = tables_per_chapter
                
            if i < total_images % len(chapters):
                chapter_image_counts[i] = images_per_chapter + 1
            else:
                chapter_image_counts[i] = images_per_chapter
    
    # Use the calculated counts
    table_counts = chapter_table_counts
    image_counts = chapter_image_counts
    
    return {
        "chapters": chapters,
        "word_counts": word_counts,
        "equation_counts": equation_counts,
        "table_counts": table_counts,
        "image_counts": image_counts,
        "paragraph_counts": paragraph_counts
    }


def get_ref_stats(docx_path: str) -> Dict[str, Any]:
    """Extract reference statistics"""
    doc = Document(docx_path)
    text_content = _get_full_text(doc)
    
    # Find references section - look for the section after 参考文献
    ref_pattern = r'参考文献\s*\n(.+?)(?=\n\s*\n|致谢|附录|\Z)'
    ref_match = re.search(ref_pattern, text_content, re.DOTALL)
    
    if not ref_match:
        # Try alternative pattern
        ref_pattern = r'参考文献(.+?)(?=\n\s*\n|致谢|附录|\Z)'
        ref_match = re.search(ref_pattern, text_content, re.DOTALL)
    
    if not ref_match:
        return {
            "total_references": 0,
            "by_indicator": {
                "期刊论文[J]": 0,
                "会议论文[C]": 0,
                "学位论文[D]": 0,
                "技术报告[R]": 0,
                "其他": 0
            },
            "by_lang": {"中文文献": 0, "英文文献": 0},
            "recent_3y": 0
        }
    
    ref_text = ref_match.group(1)
    
    # Count total references by looking for numbered reference patterns
    ref_patterns = [
        r'\[\d+\]',  # [1], [2], etc.
        r'^\d+\.',   # 1., 2., etc.
        r'^\d+\s',   # 1 , 2 , etc.
    ]
    
    total_references = 0
    for pattern in ref_patterns:
        matches = re.findall(pattern, ref_text, re.MULTILINE)
        if matches:
            total_references = max(total_references, len(matches))
    
    # Find all reference entries
    ref_lines = [line.strip() for line in ref_text.split('\n') if line.strip()]
    reference_entries = []
    
    # Check for numbered format first
    for line in ref_lines:
        if re.search(r'\[\d+\]', line):  # This is a reference line
            reference_entries.append(line)
    
    # If no numbered format found, treat each non-empty line as a reference
    if not reference_entries:
        reference_entries = [line for line in ref_lines if line.strip()]
        total_references = len(reference_entries)
    
    # Count by indicator
    indicators = {
        "期刊论文[J]": len(re.findall(r'\[J\]', ref_text)),
        "会议论文[C]": len(re.findall(r'\[C\]', ref_text)),
        "学位论文[D]": len(re.findall(r'\[D\]', ref_text)),
        "技术报告[R]": len(re.findall(r'\[R\]', ref_text)),
        "其他": 0
    }
    
    # Calculate "其他"
    total_categorized = sum(indicators.values())
    indicators["其他"] = max(0, total_references - total_categorized)
    
    # Count language for each reference entry based on first word
    chinese_refs = 0
    english_refs = 0
    
    for ref_entry in reference_entries:
        # Remove the reference number [1], [2], etc.
        content = re.sub(r'^\[\d+\]', '', ref_entry).strip()
        
        # Extract the first word (author name)
        first_word_match = re.match(r'^([^\s,，.]+)', content)
        if first_word_match:
            first_word = first_word_match.group(1)
            
            # Check if first word contains Chinese characters
            if re.search(r'[\u4e00-\u9fff]', first_word):
                chinese_refs += 1
            else:
                english_refs += 1
        else:
            # If no first word found, assume English
            english_refs += 1
    
    # Count recent 3 years references (2023-2025)
    recent_3y = 0
    for year in range(2023, 2026):  # 2023, 2024, 2025
        recent_3y += len(re.findall(str(year), ref_text))
    
    # Ensure indicators sum equals total_references
    indicator_sum = sum(indicators.values())
    if indicator_sum != total_references:
        diff = total_references - indicator_sum
        indicators["其他"] = max(0, indicators["其他"] + diff)
    
    return {
        "total_references": total_references,
        "by_indicator": indicators,
        "by_lang": {
            "中文文献": chinese_refs,
            "英文文献": english_refs
        },
        "recent_3y": recent_3y
    }


def main():
    """Example usage of the four functions"""
    import sys
    
    if len(sys.argv) != 2:
        print("Usage: python docx_analysis_functions.py <docx_file_path>")
        sys.exit(1)
    
    docx_path = sys.argv[1]
    
    try:
        print("Basic Info:")
        basic_info = get_basic_info(docx_path)
        print(json.dumps(basic_info, ensure_ascii=False, indent=2))
        
        print("\nOverall Stats:")
        overall_stats = get_overall_stats(docx_path)
        print(json.dumps(overall_stats, ensure_ascii=False, indent=2))
        
        print("\nChapter Stats:")
        chapter_stats = get_chapter_stats(docx_path)
        print(json.dumps(chapter_stats, ensure_ascii=False, indent=2))
        
        print("\nReference Stats:")
        ref_stats = get_ref_stats(docx_path)
        print(json.dumps(ref_stats, ensure_ascii=False, indent=2))
        
    except Exception as e:
        print(f"Error processing {docx_path}: {e}")
        sys.exit(1)


if __name__ == "__main__":
    main()