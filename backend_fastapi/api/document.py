"""
文档处理API路由
提供文档上传、转换、预处理等功能
"""

import os
import sys
import uuid
import base64
import tempfile
import asyncio
import re
from pathlib import Path
from datetime import datetime
from typing import Dict, Any, Optional, Tuple, List
from typing import Dict, Any, Optional
from fastapi import APIRouter, File, UploadFile, HTTPException, Form
from fastapi.responses import JSONResponse, Response
from pydantic import BaseModel

sys.path.insert(0, str(Path(__file__).parent.parent))

from config.data_config import FILE_CONFIG
from tools.docx_tools.docx2md import convert_docx_bytes_to_md
from tools.docx_tools.md2pkl import convert_md_content_to_pkl_data
from tools.docx_tools.pkl_analyse import analyze_pkl_structure, get_document_summary
from tools.logger import get_logger
from utils.redis_client import get_redis_manager

logger = get_logger(__name__)
router = APIRouter()

# 响应模型
class DocumentUploadResponse(BaseModel):
    """文档上传响应模型"""
    task_id: str
    filename: str
    file_size: int
    upload_time: str
    status: str

class DocumentProcessResponse(BaseModel):
    """文档处理响应模型"""
    success: bool
    message: str
    document_id: str
    markdown_content: Optional[str] = None
    images: Optional[list] = None
    references: Optional[List[str]] = None
    structure: Optional[Dict[str, Any]] = None
    summary: Optional[Dict[str, Any]] = None

class DocumentPreviewResponse(BaseModel):
    """文档预览响应模型"""
    success: bool
    message: str
    document_id: str
    content: Optional[str] = None
    images: Optional[list] = None
    references: Optional[List[str]] = None
    chapters: Optional[list] = None

# Redis存储（替代内存存储）
# document_storage = {}  # 已弃用，使用Redis存储

def validate_file(file: UploadFile) -> tuple[bool, str]:
    """
    验证上传的文件
    
    Args:
        file: 上传的文件
        
    Returns:
        tuple[bool, str]: (是否有效, 错误信息)
    """
    # 检查文件扩展名
    if not file.filename:
        return False, "文件名不能为空"
    
    file_ext = os.path.splitext(file.filename)[1].lower()
    if file_ext not in FILE_CONFIG['allowed_extensions']:
        return False, f"不支持的文件格式: {file_ext}。支持的格式: {', '.join(FILE_CONFIG['allowed_extensions'])}"
    
    # 检查文件大小（如果可以获取）
    if hasattr(file, 'size') and file.size:
        if file.size > FILE_CONFIG['max_file_size']:
            return False, f"文件大小超过限制: {file.size} > {FILE_CONFIG['max_file_size']}"
    
    return True, ""

@router.post("/upload", response_model=DocumentUploadResponse)
async def upload_document(file: UploadFile = File(...)):
    """
    上传文档文件
    
    Args:
        file: 上传的文档文件
        
    Returns:
        DocumentUploadResponse: 上传结果
    """
    try:
        logger.info(f"开始上传文档: {file.filename}")
        
        # 验证文件
        is_valid, error_msg = validate_file(file)
        if not is_valid:
            logger.warning(f"文件验证失败: {error_msg}")
            raise HTTPException(status_code=400, detail=error_msg)
        
        # 读取文件内容
        file_content = await file.read()
        file_size = len(file_content)
        
        # 再次检查文件大小
        if file_size > FILE_CONFIG['max_file_size']:
            raise HTTPException(
                status_code=413, 
                detail=f"文件大小超过限制: {file_size} > {FILE_CONFIG['max_file_size']}"
            )
        
        # 生成文档ID
        document_id = str(uuid.uuid4())
        
        # 存储文档信息
        document_info = {
            'id': document_id,
            'filename': file.filename,
            'content_type': file.content_type,
            'size': file_size,
            'content': file_content,
            'status': 'uploaded'
        }
        
        # 使用Redis存储文档信息
        redis_mgr = await get_redis_manager()
        if not await redis_mgr.store_document(document_id, document_info):
            raise HTTPException(status_code=500, detail="文档存储失败")
        
        logger.info(f"文档上传成功: {file.filename}, ID: {document_id}")
        
        # 按照前端API需求返回格式
        return {
            "task_id": document_id,
            "filename": file.filename,
            "file_size": file_size,
            "upload_time": datetime.now().isoformat(),
            "status": "uploaded"
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"文档上传失败: {e}")
        raise HTTPException(status_code=500, detail=f"文档上传失败: {str(e)}")

# 新的处理接口，符合前端API需求
@router.post("/process")
async def process_document_api(request: dict):
    """
    开始分析已上传的文档

    Args:
        request: 包含task_id和model_config的请求体

    Returns:
        dict: 处理结果
    """
    task_id = request.get('task_id')
    model_config = request.get('model_config', {})
    redis_mgr = None # 初始化为None

    try:
        if not task_id:
            raise HTTPException(status_code=400, detail="缺少task_id参数")

        logger.info(f"开始处理文档: {task_id}")

        # 获取Redis管理器
        redis_mgr = await get_redis_manager()

        # 尝试获取分布式锁
        lock_acquired = await redis_mgr.acquire_lock(task_id)
        if not lock_acquired:
            raise HTTPException(status_code=409, detail="文档正在处理中，请稍后再试。")

        # 检查文档是否存在
        document_info = await redis_mgr.get_document(task_id)
        if document_info is None:
            raise HTTPException(status_code=404, detail="文档不存在")

        # 再次检查文档状态（在获取锁后再次确认，防止竞态条件）
        if document_info['status'] == 'processing':
            # 如果在获取锁之前，另一个进程已经将状态设置为processing，则释放锁并报错
            await redis_mgr.release_lock(task_id)
            raise HTTPException(status_code=409, detail="文档正在处理中")

        # 更新状态
        document_info['status'] = 'processing'
        await redis_mgr.store_document(task_id, document_info)

        # 转换为Markdown
        logger.info(f"转换文档为Markdown: {task_id}")
        md_content = await asyncio.to_thread(convert_docx_bytes_to_md, document_info['content'])

        # 提取图片信息
        logger.info(f"提取文档图片: {task_id}")
        images = await asyncio.to_thread(extract_images_from_docx, document_info['content'])
        logger.info(f"提取到 {len(images)} 张图片")

        # 提取参考文献
        logger.info(f"提取参考文献: {task_id}")
        references = await asyncio.to_thread(extract_references_from_markdown, md_content)

        # 转换为结构化数据
        logger.info(f"转换为结构化数据: {task_id}")
        pkl_data = await asyncio.to_thread(convert_md_content_to_pkl_data, md_content)

        # 更新进度：基础处理完成
        document_info['md_content'] = md_content
        document_info['images'] = images
        document_info['references'] = references
        document_info['pkl_data'] = pkl_data
        document_info['progress'] = 0.3
        document_info['message'] = '基础处理完成，开始分析评估...'
        await redis_mgr.store_document(task_id, document_info)

        # 分析文档结构
        logger.info(f"分析文档结构: {task_id}")
        temp_file_path = None
        try:
            with tempfile.NamedTemporaryFile(suffix='.pkl', delete=False) as temp_file:
                import pickle
                await asyncio.to_thread(pickle.dump, pkl_data, temp_file)
                temp_file.flush()
                temp_file_path = temp_file.name

            structure = await asyncio.to_thread(analyze_pkl_structure, temp_file_path)
            summary = await asyncio.to_thread(get_document_summary, temp_file_path)

        finally:
            # 安全删除临时文件
            if temp_file_path and os.path.exists(temp_file_path):
                max_retries = 3
                for attempt in range(max_retries):
                    try:
                        await asyncio.to_thread(os.unlink, temp_file_path)
                        break
                    except (PermissionError, OSError) as e:
                        if attempt < max_retries - 1:
                            import time
                            await asyncio.to_thread(time.sleep, 0.1 * (attempt + 1))
                        else:
                            logger.warning(f"无法删除临时文件 {temp_file_path}: {e}")

        document_info['structure'] = structure
        document_info['summary'] = summary
        document_info['progress'] = 0.5
        document_info['message'] = '文档结构分析完成，标记为已处理状态...'
        # 在执行评估之前，先将状态设置为processed，确保评估函数能获取到markdown内容
        document_info['status'] = 'processed'
        document_info['processed_at'] = datetime.now().isoformat()
        await redis_mgr.store_document(task_id, document_info)

        # 启动后台评估任务
        logger.info(f"文档处理完成，启动后台评估任务: {task_id}")
        try:
            from utils.async_tasks import get_task_manager
            task_manager = await get_task_manager()
            await task_manager.start_evaluation_task(task_id, document_info)
            document_info['progress'] = 0.6
            document_info['message'] = '文档处理完成，正在进行后台分析评估...'
            logger.info(f"后台评估任务启动成功: {task_id}")
        except Exception as e:
            logger.error(f"启动后台评估任务失败: {task_id}, 错误: {e}")
            import traceback
            logger.error(f"详细错误信息: {traceback.format_exc()}")
            document_info['progress'] = 0.9
            document_info['message'] = '文档处理完成，可通过analysis接口进行评估...'

        await redis_mgr.store_document(task_id, document_info)

        # 生成HTML预览文件 - 使用docx2html.py
        logger.info(f"生成HTML预览文件: {task_id}")
        try:
            # 创建预览目录
            preview_dir = Path("preview") / task_id
            preview_dir.mkdir(parents=True, exist_ok=True)
            images_dir = preview_dir / "images"
            images_dir.mkdir(exist_ok=True)

            # 保存原始DOCX文件到临时位置
            temp_docx_path = preview_dir / "temp_document.docx"
            await asyncio.to_thread(lambda: open(temp_docx_path, 'wb').write(document_info['content']))

            # 使用docx2html.py转换文档
            from tools.docx_tools.docx2html import docx_to_html
            html_file_path = preview_dir / "document.html"

            # 转换DOCX到HTML，图片会自动保存到images目录
            await asyncio.to_thread(
                docx_to_html,
                docx_path=str(temp_docx_path),
                output_html_path=str(html_file_path),
                image_dir=str(images_dir)
            )

            # 清理临时DOCX文件
            await asyncio.to_thread(os.unlink, temp_docx_path)

            logger.info(f"HTML预览文件生成成功: {html_file_path}")

            # 读取生成的HTML内容以便后续处理
            html_content = await asyncio.to_thread(lambda: open(html_file_path, 'r', encoding='utf-8').read())

            # 备用：如果docx2html生成的图片不够，补充保存从文档提取的图片
            if images:
                logger.info(f"补充保存 {len(images)} 张提取的图片到 {images_dir}")
                import base64

                for i, img_info in enumerate(images):
                    try:
                        # 使用1基索引保存图片文件，与前端请求保持一致
                        img_path = images_dir / f"image_{i+1}.png"

                        # 如果docx2html已经生成了这个图片，跳过
                        if await asyncio.to_thread(os.path.exists, img_path):
                            logger.info(f"图片 {i+1} 已存在，跳过: {img_path}")
                            continue

                        # img_info是字典，包含image_data(base64编码)
                        if isinstance(img_info, dict) and 'image_data' in img_info:
                            # 解码base64数据
                            img_bytes = await asyncio.to_thread(base64.b64decode, img_info['image_data'])
                            await asyncio.to_thread(lambda: open(img_path, 'wb').write(img_bytes))
                            logger.info(f"补充保存图片 {i+1}: {img_path} (大小: {len(img_bytes)} bytes)")
                        elif isinstance(img_info, bytes):
                            # 兼容旧格式
                            await asyncio.to_thread(lambda: open(img_path, 'wb').write(img_info))
                            logger.info(f"补充保存图片 {i+1}: {img_path} (旧格式)")
                        else:
                            logger.warning(f"图片 {i+1} 格式不正确: {type(img_info)}")

                    except Exception as e:
                        logger.error(f"保存图片 {i+1} 失败: {e}")

            # 将HTML文件路径保存到文档信息中
            document_info['html_file_path'] = str(html_file_path)
        except Exception as e:
            logger.warning(f"生成HTML预览文件失败: {e}")
            document_info['html_generation_error'] = str(e)

        # 标记为完全处理完成
        document_info['progress'] = 1.0
        document_info['message'] = '所有分析完成！'
        await redis_mgr.store_document(task_id, document_info)

        logger.info(f"文档处理完成: {task_id}")

        return {
            "task_id": task_id,
            "status": "processing",
            "message": "开始处理文档"
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"文档处理失败: {e}")
        # 更新状态为失败
        if redis_mgr:
            try:
                document_info = await redis_mgr.get_document(task_id)
                if document_info:
                    document_info['status'] = 'failed'
                    document_info['error'] = str(e)
                    await redis_mgr.store_document(task_id, document_info)
            except Exception as redis_error:
                logger.error(f"更新文档失败状态失败: {redis_error}")
        raise HTTPException(status_code=500, detail=f"文档处理失败: {str(e)}")
    finally:
        if redis_mgr:
            await redis_mgr.release_lock(task_id)

@router.post("/process/{document_id}", response_model=DocumentProcessResponse)
async def process_document(document_id: str):
    """
    处理已上传的文档，转换为结构化数据
    
    Args:
        document_id: 文档ID
        
    Returns:
        DocumentProcessResponse: 处理结果
    """
    try:
        logger.info(f"开始处理文档: {document_id}")
        
        # 获取Redis管理器
        redis_mgr = await get_redis_manager()
        
        # 检查文档是否存在
        document_info = await redis_mgr.get_document(document_id)
        if document_info is None:
            raise HTTPException(status_code=404, detail="文档不存在")
        
        # 检查文档状态
        if document_info['status'] == 'processing':
            raise HTTPException(status_code=409, detail="文档正在处理中")
        
        # 更新状态
        document_info['status'] = 'processing'
        await redis_mgr.store_document(document_id, document_info)
        
        # 转换为Markdown
        logger.info(f"转换文档为Markdown: {document_id}")
        md_content = await asyncio.to_thread(convert_docx_bytes_to_md, document_info['content'])
        
        # 提取图片信息
        logger.info(f"提取文档图片: {document_id}")
        images = await asyncio.to_thread(extract_images_from_docx, document_info['content'])
        logger.info(f"提取到 {len(images)} 张图片")
        
        # 提取参考文献
        logger.info(f"提取参考文献: {document_id}")
        references = await asyncio.to_thread(extract_references_from_markdown, md_content)
        
        # 转换为结构化数据
        logger.info(f"转换为结构化数据: {document_id}")
        pkl_data = await asyncio.to_thread(convert_md_content_to_pkl_data, md_content)
        
        # 保存处理结果到文档信息
        document_info['md_content'] = md_content
        document_info['images'] = images
        document_info['references'] = references
        document_info['pkl_data'] = pkl_data
        document_info['status'] = 'processed'
        
        # 分析文档结构
        # 创建临时文件用于分析
        temp_file_path = None
        try:
            with tempfile.NamedTemporaryFile(suffix='.pkl', delete=False) as temp_file:
                import pickle
                await asyncio.to_thread(pickle.dump, pkl_data, temp_file)
                temp_file.flush()
                temp_file_path = temp_file.name

            structure = await asyncio.to_thread(analyze_pkl_structure, temp_file_path)
            summary = await asyncio.to_thread(get_document_summary, temp_file_path)

        finally:
            # 安全删除临时文件
            if temp_file_path and os.path.exists(temp_file_path):
                max_retries = 3
                for attempt in range(max_retries):
                    try:
                        await asyncio.to_thread(os.unlink, temp_file_path)
                        break
                    except (PermissionError, OSError) as e:
                        if attempt < max_retries - 1:
                            import time
                            await asyncio.to_thread(time.sleep, 0.1 * (attempt + 1))
                        else:
                            logger.warning(f"无法删除临时文件 {temp_file_path}: {e}")
        
        document_info['structure'] = structure
        document_info['summary'] = summary
        
        # 保存到Redis
        await redis_mgr.store_document(document_id, document_info)
        
        logger.info(f"文档处理完成: {document_id}")
        
        return DocumentProcessResponse(
            success=True,
            message="文档处理完成",
            document_id=document_id,
            markdown_content=md_content,
            images=images,
            references=references,
            structure=structure,
            summary=summary
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"文档处理失败: {e}")
        # 更新状态为失败
        try:
            redis_mgr = await get_redis_manager()
            document_info = await redis_mgr.get_document(document_id)
            if document_info:
                document_info['status'] = 'failed'
                document_info['error'] = str(e)
                await redis_mgr.store_document(document_id, document_info)
        except Exception as redis_error:
            logger.error(f"更新文档失败状态失败: {redis_error}")
        raise HTTPException(status_code=500, detail=f"文档处理失败: {str(e)}")

@router.get("/preview/{document_id}", response_model=DocumentPreviewResponse)
async def preview_document(document_id: str):
    """
    预览文档内容
    
    Args:
        document_id: 文档ID
        
    Returns:
        DocumentPreviewResponse: 预览结果
    """
    try:
        logger.info(f"预览文档: {document_id}")
        
        # 获取Redis管理器
        redis_mgr = await get_redis_manager()
        
        # 检查文档是否存在
        document_info = await redis_mgr.get_document(document_id)
        if document_info is None:
            raise HTTPException(status_code=404, detail="文档不存在")
        
        # 检查文档是否已处理
        if document_info['status'] != 'processed':
            raise HTTPException(status_code=400, detail="文档尚未处理完成")
        
        # 获取markdown内容和图片
        md_content = document_info.get('md_content', '')
        images = document_info.get('images', [])
        references = document_info.get('references', [])
        
        # 获取章节信息
        pkl_data = document_info.get('pkl_data', {})
        chapters = pkl_data.get('chapters', [])
        
        # 格式化章节信息用于预览
        chapter_list = []
        for i, chapter in enumerate(chapters):
            chapter_list.append({
                'index': i,
                'name': chapter.get('chapter_name', f'第{i+1}章'),
                'content_length': len(chapter.get('content', '')),
                'image_count': len(chapter.get('images', []))
            })
        
        return DocumentPreviewResponse(
            success=True,
            message="文档预览获取成功",
            document_id=document_id,
            content=md_content,
            images=images,
            references=references,
            chapters=chapter_list
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"文档预览失败: {e}")
        raise HTTPException(status_code=500, detail=f"文档预览失败: {str(e)}")



async def get_document_markdown(document_id: str) -> Optional[str]:
    """
    从Redis中获取指定文档的markdown内容
    
    Args:
        document_id: 文档ID
        
    Returns:
        Optional[str]: markdown内容，如果文档不存在或未处理则返回None
    """
    try:
        logger.info(f"获取文档 {document_id} 的Markdown内容")
        # 获取Redis管理器
        redis_mgr = await get_redis_manager()
        
        # 获取文档信息
        document_info = await redis_mgr.get_document(document_id)
        if document_info is None:
            logger.warning(f"文档不存在: {document_id}")
            return None
        
        # 检查文档是否已处理
        if document_info['status'] != 'processed':
            logger.warning(f"文档尚未处理完成: {document_id}, status: {document_info['status']}")
            return None
        
        # 返回markdown内容
        md_content = document_info.get('md_content', '')
        if not md_content:
            logger.warning(f"文档没有Markdown内容: {document_id}")
            return None
        logger.info(f"成功获取文档markdown内容: {document_id}, 长度: {len(md_content)}")
        return md_content
    
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"获取Markdown内容失败: {e}")
        raise HTTPException(status_code=500, detail=f"获取Markdown内容失败: {str(e)}")

def extract_references_from_markdown(md_content: str) -> List[str]:
    """
    从Markdown内容中提取参考文献

    Args:
        md_content: Markdown内容

    Returns:
        List[str]: 参考文献列表，每条参考文献删除换行符
    """
    references = []

    # 多种参考文献章节标题模式
    ref_patterns = [
        r'#\s*参考文献',
        r'#\s*References',
        r'#\s*REFERENCES',
        r'##\s*参考文献',
        r'##\s*References',
        r'###\s*参考文献',
        r'参考文献',
        r'References'
    ]

    ref_match = None
    for pattern in ref_patterns:
        ref_match = re.search(pattern, md_content, re.IGNORECASE)
        if ref_match:
            logger.info(f"找到参考文献章节，使用模式: {pattern}")
            break

    if not ref_match:
        logger.warning("未找到参考文献章节，尝试在整个文档中查找参考文献条目")
        # 如果没有找到参考文献章节，尝试在整个文档中查找参考文献条目
        ref_content = md_content
    else:
        # 获取参考文献章节开始位置
        ref_start = ref_match.end()

        # 查找下一个章节（如致谢、附录等）的开始位置作为结束位置
        # 匹配以#开头的章节标题
        next_section_pattern = r'\n#+\s+(?!参考文献|References|REFERENCES)'
        next_section_match = re.search(next_section_pattern, md_content[ref_start:], re.IGNORECASE)

        if next_section_match:
            ref_end = ref_start + next_section_match.start()
            ref_content = md_content[ref_start:ref_end]
        else:
            # 如果没有找到下一个章节，则参考文献内容到文档结尾
            ref_content = md_content[ref_start:]

    # 多种参考文献条目模式
    ref_item_patterns = [
        r'\[(\d+)\]\s*(.+?)(?=(?:\n\s*\[\d+\]|\n#+|\\Z))',  # [1] 格式，支持多行
        r'^\s*\[(\d+)\]\s*(.+?)(?=(?:\n\s*\[\d+\]|\n#+|\\Z))',  # 行首[1] 格式，支持多行
        r'^\s*(\d+)\.\s*(.+?)(?=(?:\n\s*\d+\.|\n#+|$))',  # 行首1. 格式，支持多行
        r'^\s*(\d+)\)\s*(.+?)(?=(?:\n\s*\d+\)|\n#+|$))',  # 行首1) 格式，支持多行
    ]

    # 先尝试带编号的格式
    found_numbered_refs = False
    for pattern in ref_item_patterns:
        ref_matches = re.findall(pattern, ref_content, re.DOTALL | re.MULTILINE)
        if ref_matches:
            # 过滤掉年份等错误匹配（真正的参考文献编号应该是1, 2, 3...这样的小数字）
            valid_refs = []
            for ref_num, ref_text in ref_matches:
                # 参考文献编号应该是合理的范围（1-1000），不是年份
                if int(ref_num) <= 1000:
                    clean_ref = re.sub(r'\s+', ' ', ref_text.strip())
                    if clean_ref and len(clean_ref) > 10:  # 过滤太短的条目
                        valid_refs.append((ref_num, clean_ref))
            
            if valid_refs:
                logger.info(f"使用带编号模式提取参考文献: {pattern}")
                for ref_num, clean_ref in valid_refs:
                    if pattern.startswith(r'\[') or r'\[' in pattern:
                        references.append(f"[{ref_num}] {clean_ref}")
                    else:
                        references.append(f"{ref_num}. {clean_ref}")
                found_numbered_refs = True
                break
    
    # 如果没有找到带编号的格式，尝试无编号格式（每行一个条目）
    if not found_numbered_refs:
        logger.info("未找到带编号的参考文献，尝试无编号格式")
        
        # 按行分割，过滤空行
        lines = ref_content.split('\n')
        non_empty_lines = [line.strip() for line in lines if line.strip()]
        
        # 检查是否是参考文献格式（包含作者名、期刊等特征）
        ref_patterns_check = [
            r'et al\.',  # 等人
            r'\[J\]',    # 期刊
            r'\[C\]',    # 会议
            r'arXiv',    # arXiv预印本
            r'Proceedings',  # 会议论文集
            r'\..*\d{4}',    # 包含年份
        ]
        
        for line in non_empty_lines:
            # 检查是否包含参考文献的特征
            is_reference = any(re.search(pattern, line, re.IGNORECASE) for pattern in ref_patterns_check)
            # 或者检查是否包含常见的作者名格式（大写字母开头，包含逗号分隔）
            is_author_format = re.search(r'^[A-Z][a-z]+\s+[A-Z]', line)
            
            if is_reference or is_author_format:
                # 删除多余的空白字符，但保留基本格式
                clean_ref = re.sub(r'\s+', ' ', line.strip())
                if clean_ref and len(clean_ref) > 20:  # 参考文献通常比较长
                    references.append(clean_ref)
        
        if references:
            logger.info(f"使用无编号格式提取到 {len(references)} 条参考文献")

    logger.info(f"提取到 {len(references)} 条参考文献")
    return references


def extract_images_from_docx(docx_bytes: bytes) -> List[Dict[str, Any]]:
    """
    从DOCX文件中提取图片及其标题

    Args:
        docx_bytes: DOCX文件的字节数据

    Returns:
        List[Dict[str, Any]]: 图片信息列表，包含image_id、image_data、image_type、image_title等
    """
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from io import BytesIO

        # 使用临时文件处理docx，在Windows上需要设置delete=False
        temp_file_path = None
        try:
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
                temp_file.write(docx_bytes)
                temp_file.flush()
                temp_file_path = temp_file.name

            # 在临时文件关闭后再打开文档
            doc = Document(temp_file_path)
            images = []

            # 先构建图片ID到文件名的映射
            image_rels = {}
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    image_rels[rel.target_ref] = rel.target_part.blob

            # 遍历文档段落，查找图片和对应的标题
            for para_idx, paragraph in enumerate(doc.paragraphs):
                # 检查段落中是否包含图片
                for run in paragraph.runs:
                    # 检查run中是否有图片
                    for drawing in run.element.xpath('.//w:drawing'):
                        # 查找图片引用
                        for blip in drawing.xpath('.//a:blip'):
                            embed_id = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                            if embed_id:
                                # 查找对应的图片数据
                                rel = doc.part.rels.get(embed_id)
                                if rel and "image" in rel.target_ref:
                                    try:
                                        # 获取图片数据
                                        image_bytes = rel.target_part.blob
                                        image_ext = os.path.splitext(rel.target_ref)[-1] or '.png'

                                        # 生成图片ID
                                        image_id = f"img_{len(images) + 1}"

                                        # 将图片数据编码为base64
                                        image_base64 = base64.b64encode(image_bytes).decode('utf-8')

                                        # 查找图片标题
                                        image_title = find_image_title(doc, para_idx, image_id)

                                        images.append({
                                            'image_id': image_id,
                                            'image_data': image_base64,
                                            'image_type': image_ext.lstrip('.'),
                                            'image_title': image_title,
                                            'size': len(image_bytes)
                                        })
                                    except Exception as e:
                                        logger.warning(f"提取图片失败: {e}")
                                        continue

            # 如果没有找到图片，使用原来的方法作为备选
            if not images:
                logger.info("使用备选方法提取图片")
                for rel in doc.part.rels.values():
                    if "image" in rel.target_ref:
                        try:
                            # 获取图片数据
                            image_bytes = rel.target_part.blob
                            image_ext = os.path.splitext(rel.target_ref)[-1] or '.png'

                            # 生成图片ID
                            image_id = f"img_{len(images) + 1}"

                            # 将图片数据编码为base64
                            image_base64 = base64.b64encode(image_bytes).decode('utf-8')

                            images.append({
                                'image_id': image_id,
                                'image_data': image_base64,
                                'image_type': image_ext.lstrip('.'),
                                'image_title': image_id,  # 使用image_id作为默认标题
                                'size': len(image_bytes)
                            })
                        except Exception as e:
                            logger.warning(f"提取图片失败: {e}")
                            continue

            return images

        finally:
            # 清理临时文件
            if temp_file_path and os.path.exists(temp_file_path):
                try:
                    os.unlink(temp_file_path)
                except Exception as e:
                    logger.warning(f"无法删除临时文件 {temp_file_path}: {e}")
            
    except Exception as e:
        logger.error(f"提取图片时发生错误: {e}")
        return []


def find_image_title(doc, image_para_idx: int, default_title: str) -> str:
    """
    查找图片对应的标题
    
    Args:
        doc: Document对象
        image_para_idx: 包含图片的段落索引
        default_title: 默认标题（通常是image_id）
        
    Returns:
        str: 图片标题
    """
    try:
        # 优先检查图片紧邻的段落（后一个，前一个）
        search_offsets = [(1, 4), (-1, -3)]  # (start_offset, end_offset) for after and before
        
        for start_offset, end_offset in search_offsets:
            step = 1 if start_offset > 0 else -1
            for offset in range(start_offset, end_offset, step):
                target_para_idx = image_para_idx + offset
                
                if not (0 <= target_para_idx < len(doc.paragraphs)):
                    continue
                
                para = doc.paragraphs[target_para_idx]
                para_text = para.text.strip()
                
                if not para_text:
                    continue
                
                # 检查是否是图片标题的模式
                # 更加灵活的匹配 "图" 后面的数字和点号，允许更多空格
                title_pattern = r'^[图表]\s*\d+(\.\d+)*\s*([\.\-—_]?\s*.*)?'

                if re.match(title_pattern, para_text):
                    if is_likely_image_title(para):
                        clean_title = para_text.replace(' ', '_')
                        logger.info(f"找到图片标题: {para_text} -> {clean_title}")
                        return clean_title
                
                # 如果遇到非空的普通段落（很长的文本），停止当前方向的搜索
                if len(para_text) > 100:
                    break
        
        logger.info(f"未找到图片标题，使用默认值: {default_title}")
        return default_title
        
    except Exception as e:
        logger.warning(f"查找图片标题时发生错误: {e}")
        return default_title


def is_likely_image_title(paragraph) -> bool:
    """
    判断段落是否可能是图片标题
    
    Args:
        paragraph: 段落对象
        
    Returns:
        bool: 是否可能是图片标题
    """
    try:
        # 检查段落对齐方式（图片标题通常居中）
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        if paragraph.alignment and paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER:
            return True
        
        # 检查字体样式（楷体等）
        for run in paragraph.runs:
            if run.font.name:
                font_name = run.font.name.lower()
                # 检查是否是楷体
                if '楷' in font_name or 'kai' in font_name:
                    return True
            
            # 检查是否是加粗或斜体（图片标题可能有特殊格式）
            if run.bold or run.italic:
                return True
        
        # 如果段落文本较短且包含"图"字，也认为可能是标题
        text = paragraph.text.strip()
        if len(text) < 50 and '图' in text:
            return True
        
        return False
        
    except Exception as e:
        logger.warning(f"检查段落格式时发生错误: {e}")
        return False
